from docx import Document
from docx2pdf import convert
import os



class BuildingTheDocument:

    def __init__(self, table_data,data,table_data2):
        self.table_data = table_data
        self.data = data
        self.table_data2 = table_data2
    
    def MontandoODocx(self):

        doc = Document()
        dir = os.path.join(os.getcwd(),r"Output")
        
        doc.add_heading("Seu Guia de Investimento")

        print(self.data[1])
        for i in range(len(self.data)):
            doc.add_paragraph(self.table_data[i])
            doc.add_paragraph(self.data[i])
            doc.add_paragraph("Disclaimer")
            doc.add_paragraph(f"Essa é uma sugestão de investimento para a classe de renda {self.table_data2[i]}, e não uma recomendação exclusiva. Existem outros investimentos que podem desempenhar um papel semelhante em sua carteira, dependendo de suas necessidades e perfil de investidor.")
            
        doc.add_paragraph()
        outputDir = os.path.join(os.getcwd(),"/Output")
        doc.save(dir+r"\result.docx")
        convert(dir+r"\result.docx")
        convert(dir,r"\result.pdf")
        convert(outputDir)
