import os
from docx import Document
from docx2pdf import convert


class TakingData:
    

    def __init__(self, file_name):
        self.file_name = os.path.join(os.getcwd(),f"Data\{file_name}")
        self.table_data = []
        self.table_data2 = []
        self.data = []
        


    def FetchingTableData(self):
        doc = Document(self.file_name)
        table = doc.tables[0]
        # self.table_data = []
        # self.data = []
        first_cells_data = []
        second_cells_data = []
        first_cells_data_pop_fundos = []
        second_cells_data_pop_fundos = []

        # Pegando as informações da tabela

        for row in table.rows:
            first_cells_data.append(row.cells[0])
            second_cells_data.append(row.cells[1])
            

        for i in range(len(first_cells_data)):
            if(i == 0):
                continue
            first_cells_data_pop_fundos.append(first_cells_data[i].text)
            second_cells_data_pop_fundos.append(second_cells_data[i].text)
        
        self.table_data = first_cells_data_pop_fundos
        self.table_data2 = second_cells_data_pop_fundos

        # Pegando todos os paragrafos com informções e separando em um arquivo txt para trabalhar melhor com os dados
        go = False
        allParagraphs = []
        allLines = []

        for i in range(len(doc.paragraphs)-1):
            if(go):
                allParagraphs.append(doc.paragraphs[i].text)
            if("Lista de Trends" in doc.paragraphs[i].text): # só começa a adicionar os paragrafos depois da Lista de Trends para não trazer muita informção desnecessária
                go = True

        with open('WritenContent/textData.txt', 'w') as arquivo:
            for elemento in allParagraphs:
                arquivo.write(str(elemento) + '\n')

        # Eliminando as linhas vazias 
        with open('WritenContent/textData.txt', 'r') as arquivo:
            for line in arquivo:
                if(line.strip() == "" or line.strip() =="\n"):
                    continue
                allLines.append(line)

        indexesBegin = []
        indexesEnd = []

        # pegando o começo e o fim do porque está nessa lista de sugestões
        for i in range(len(allLines)):
            if (allLines[i].strip() == "Por que está nessa lista de sugestões?"):
                indexesBegin.append(i+1)
            if (allLines[i].strip() == "Confira nossa visão para essa classe de ativo aqui."):
                indexesEnd.append(i)
        # print(indexesBegin)
        # print(indexesEnd)
        
        finalLine = ""

        for i in range(0,8):
            for j in range(indexesBegin[i], indexesEnd[i]):
                finalLine += allLines[j]
            self.data.append(finalLine)
            finalLine = ""

        
            
        
        
        


        



            
            

        









    



if __name__ == "__main__":
    dd = TakingData("Guia de Investimentos (lista de Trends).docx")
    dd.FetchingTableData()